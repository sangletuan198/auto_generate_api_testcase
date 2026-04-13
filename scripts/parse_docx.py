#!/usr/bin/env python3
"""
parse_docx.py
━━━━━━━━━━━━━
Parse API specification from .docx (and legacy .doc) files.
Produces the SAME contract format as parse_html_docs.py so the rest
of the pipeline (regen_from_contracts, merge, compare) works unchanged.

Supported formats:
  - .docx (Word 2007+) — parsed directly with python-docx
  - .doc  (legacy Word 97-2003) — auto-converted to .docx first
    Conversion uses: textutil (macOS built-in) or libreoffice (Linux/Win)

Supported table layouts:
  - Confluence-exported Word docs
  - Any Word doc that has the standard tables:
      1. Overview table  → contains "API Name", "Method" rows
      2. Request table   → columns: No | Parameter | Level | Description | Type | MaxLen | Mandatory | … | Note
      3. Error table     → columns: STT | Error code | Messagekey | DeclineReason | …

Usage (standalone):
    python3 parse_docx.py              # parse all .doc/.docx in docs/
    python3 parse_docx.py path/to.docx # parse specific file
    python3 parse_docx.py path/to.doc  # auto-convert then parse

When imported:
    from parse_docx import parse_docx_file, convert_doc_to_docx
    contract = parse_docx_file(Path("docs/my_api.docx"))
"""

import platform
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Tuple, Dict, Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from soap_body_utils import detect_soap_body, soap_body_to_flat_dict
    _HAS_SOAP_UTILS = True
except ImportError:
    _HAS_SOAP_UTILS = False


# ---------------------------------------------------------------------------
# .doc → .html conversion  (tables are preserved — unlike .doc→.docx)
# ---------------------------------------------------------------------------

def convert_doc_to_html(doc_path: Path, out_dir: Optional[Path] = None) -> Optional[Path]:
    """
    Convert a legacy .doc file to .html, preserving table structure.

    Strategy (first that succeeds):
      1. textutil -convert html  — macOS built-in, fast, preserves tables
      2. libreoffice --headless --convert-to html  — cross-platform fallback

    NOTE: textutil .doc→.docx loses tables, but .doc→.html preserves them perfectly.

    Returns the Path to the converted .html, or None on failure.
    The .html is written to *out_dir* (default: same directory as source).
    """
    doc_path = Path(doc_path).resolve()
    if not doc_path.exists():
        return None

    if out_dir is None:
        out_dir = doc_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    html_out = out_dir / (doc_path.stem + '._converted.html')

    # ── 1. textutil (macOS) ───────────────────────────────────────────
    if shutil.which('textutil'):
        try:
            subprocess.run(
                ['textutil', '-convert', 'html', str(doc_path), '-output', str(html_out)],
                check=True, capture_output=True, timeout=30,
            )
            if html_out.exists() and html_out.stat().st_size > 0:
                return html_out
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            print(f'  ⚠️  textutil failed for {doc_path.name}: {exc}')

    # ── 2. libreoffice --headless ─────────────────────────────────────
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    # On Windows, LibreOffice may not be on PATH — check common install locations
    if not soffice and platform.system() == 'Windows':
        for candidate in [
            Path(r'C:\Program Files\LibreOffice\program\soffice.exe'),
            Path(r'C:\Program Files (x86)\LibreOffice\program\soffice.exe'),
        ]:
            if candidate.exists():
                soffice = str(candidate)
                break
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as tmp:
                subprocess.run(
                    [soffice, '--headless', '--convert-to', 'html',
                     '--outdir', tmp, str(doc_path)],
                    check=True, capture_output=True, timeout=60,
                )
                converted = Path(tmp) / (doc_path.stem + '.html')
                if converted.exists():
                    # shutil.move can fail on Windows if target exists
                    if html_out.exists():
                        html_out.unlink()
                    shutil.move(str(converted), str(html_out))
                    return html_out
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            print(f'  ⚠️  libreoffice failed for {doc_path.name}: {exc}')

    # ── 3. python-docx fallback (some .doc files are actually .docx renamed) ──
    if HAS_DOCX:
        try:
            doc = Document(str(doc_path))
            # If python-docx can open it, write as simple HTML with tables
            html_parts = ['<html><body>']
            for table in doc.tables:
                html_parts.append('<table border="1">')
                for row in table.rows:
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        html_parts.append(f'<td>{cell.text}</td>')
                    html_parts.append('</tr>')
                html_parts.append('</table>')
            html_parts.append('</body></html>')
            html_out.write_text('\n'.join(html_parts), encoding='utf-8')
            if html_out.exists() and html_out.stat().st_size > 100:
                print(f'  ℹ️  {doc_path.name} opened via python-docx fallback (may be .docx renamed)')
                return html_out
        except Exception:
            pass  # Not a .docx in disguise — continue to error

    import platform
    os_name = platform.system()
    print(f'  ❌  Không thể convert {doc_path.name} → .html')
    if os_name == 'Darwin':
        print(f'      textutil should be available on macOS. Check: which textutil')
    elif os_name == 'Linux':
        print(f'      Cài libreoffice: sudo apt install libreoffice  (hoặc snap install libreoffice)')
    else:
        print(f'      Cài libreoffice: https://www.libreoffice.org/download/')
    print(f'      Hoặc convert file .doc → .docx bằng Word/Google Docs rồi đặt vào docs/')
    return None


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _cell_text(cell) -> str:
    """Get clean text from a docx table cell."""
    return cell.text.strip().replace('\xa0', ' ')


def _is_struck(cell) -> bool:
    """Check if cell text has strikethrough formatting."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            rpr = run._element.find(qn('w:rPr'))
            if rpr is not None:
                strike = rpr.find(qn('w:strike'))
                dstrike = rpr.find(qn('w:dstrike'))
                if strike is not None and strike.get(qn('w:val'), 'true') != 'false':
                    return True
                if dstrike is not None and dstrike.get(qn('w:val'), 'true') != 'false':
                    return True
    return False


def _kebab_to_slug(s: str) -> str:
    """Convert /get-saving-account-transactions → getSavingAccountTransactions."""
    s = s.strip('/ ').strip()
    parts = re.split(r'[-_]+', s)
    return parts[0] + ''.join(p.capitalize() for p in parts[1:])


def _normalize_code(s: str) -> str:
    """Normalise 'DT. 005.01.000' → 'DT.005.01.000'."""
    return re.sub(r'[\s\xa0]+', '', s)


def _table_to_rows(table) -> List[List[str]]:
    """Convert a docx table to a list of rows (list of cell texts)."""
    rows = []
    for row in table.rows:
        rows.append([_cell_text(c) for c in row.cells])
    return rows


def _table_header_text(table) -> str:
    """Get concatenated text of first row cells (lowered) for matching."""
    if not table.rows:
        return ''
    return ' '.join(_cell_text(c) for c in table.rows[0].cells).lower()


def _find_table(tables, keywords: List[str]):
    """Find table whose first 2 rows contain all keywords (case-insensitive)."""
    for tbl in tables:
        sample = ''
        for row in tbl.rows[:3]:
            sample += ' '.join(_cell_text(c) for c in row.cells) + ' '
        sample = sample.lower().replace('\xa0', ' ')
        if all(kw.lower() in sample for kw in keywords):
            return tbl
    return None


# ---------------------------------------------------------------------------
# 1. Extract API Name
# ---------------------------------------------------------------------------

def extract_api_name(tables) -> str:
    """Find 'API Name' in overview table and return its value."""
    for tbl in tables:
        for row in tbl.rows:
            cells = [_cell_text(c) for c in row.cells]
            for i, text in enumerate(cells):
                if text.strip().lower() == 'api name' and i + 1 < len(cells):
                    return cells[i + 1].strip()
    return ''


# ---------------------------------------------------------------------------
# 2. Extract Method
# ---------------------------------------------------------------------------

def extract_method(tables) -> Optional[str]:
    """Find 'Method' in overview table and return its value (GET/POST/...)."""
    for tbl in tables:
        for row in tbl.rows:
            cells = [_cell_text(c) for c in row.cells]
            for i, text in enumerate(cells):
                if text.strip().lower() == 'method' and i + 1 < len(cells):
                    return cells[i + 1].strip().upper()
    return None


# ---------------------------------------------------------------------------
# 3. Parse request fields table
# ---------------------------------------------------------------------------

def _detect_col_layout(header_cells: List[str]) -> dict:
    """Auto-detect column indices from header row text."""
    h = [c.lower().replace('\xa0', ' ').strip() for c in header_cells]

    # Find key columns by header keywords
    col_map = {}
    for i, text in enumerate(h):
        if 'parameter' in text and 'name' not in col_map:
            col_map['name'] = i
        elif text in ('no', 'no.', 'stt', '#'):
            col_map['no'] = i
        elif 'level' in text:
            col_map['level'] = i
        elif 'type' in text and 'level' not in text:
            col_map['type'] = i
        elif 'mandatory' in text:
            col_map['mandatory'] = i

    # If no explicit parameter column, name is col after 'no'
    if 'name' not in col_map:
        col_map['name'] = col_map.get('no', 0) + 1

    return col_map


def parse_request_table(tables) -> Tuple[List[dict], List[dict]]:
    """Parse request body fields table. Returns (active_fields, struck_fields)."""
    tbl = (_find_table(tables, ['cifno'])
           or _find_table(tables, ['parameter', 'mandatory'])
           or _find_table(tables, ['parameter', 'type', 'level']))
    if not tbl:
        return [], []

    # Detect column layout from header row
    if not tbl.rows:
        return [], []
    header_cells = [_cell_text(c) for c in tbl.rows[0].cells]
    col_map = _detect_col_layout(header_cells)

    COL_NAME  = col_map.get('name', 1)
    COL_LEVEL = col_map.get('level', 2)
    COL_TYPE  = col_map.get('type', 4)
    COL_MAND  = col_map.get('mandatory', 6)

    fields = []
    for i, row in enumerate(tbl.rows):
        if i == 0:  # skip header
            continue
        cells = row.cells
        texts = [_cell_text(c) for c in cells]

        # Try name from COL_NAME, then COL_NAME+1 (for two-column parameter layout)
        name = texts[COL_NAME].strip('* ') if len(texts) > COL_NAME else ''
        name_col_idx = COL_NAME
        if not name and len(texts) > COL_NAME + 1:
            name = texts[COL_NAME + 1].strip('* ')
            name_col_idx = COL_NAME + 1

        # Clean name (remove whitespace artifacts)
        name = re.sub(r'\s+', '', name)
        if not name or re.match(r'^\d+$', name):
            continue

        level = texts[COL_LEVEL].strip() if len(texts) > COL_LEVEL else ''
        typ   = texts[COL_TYPE].strip() if len(texts) > COL_TYPE else ''
        mand  = texts[COL_MAND].strip() if len(texts) > COL_MAND else ''
        note  = texts[-1].strip() if texts else ''

        # Strikethrough detection on name cell
        struck = _is_struck(cells[name_col_idx]) if len(cells) > name_col_idx else False

        fields.append({
            'name': name,
            'level': level,
            'type': typ,
            'mandatory': mand,
            'note': note,
            'struck': struck,
        })

    active = [f for f in fields if not f['struck']]
    struck_out = [f for f in fields if f['struck']]
    return active, struck_out


# ---------------------------------------------------------------------------
# 4. Parse error codes table
# ---------------------------------------------------------------------------

def parse_error_table(tables) -> Tuple[List[dict], List[dict]]:
    """Parse error codes table. Returns (active_errors, struck_errors)."""
    tbl = (_find_table(tables, ['unauthorized'])
           or _find_table(tables, ['cif_not_found'])
           or _find_table(tables, ['account_not_found'])
           or _find_table(tables, ['declinreason', 'stt'])
           or _find_table(tables, ['error code', 'messagekey'])
           or _find_table(tables, ['errorcode', 'message']))
    if not tbl:
        return [], []

    errors = []
    for i, row in enumerate(tbl.rows):
        if i == 0:  # skip header
            continue
        cells = row.cells
        texts = [_cell_text(c) for c in cells]
        if not texts:
            continue

        # Error code in col 1 (col 0 = STT/sequence number)
        raw_code = texts[1] if len(texts) > 1 else ''
        code = _normalize_code(raw_code)
        if not code or not re.match(r'^[A-Z]{1,5}\.', code):
            continue

        key    = texts[2].strip() if len(texts) > 2 else ''
        reason = texts[3].strip() if len(texts) > 3 else ''

        # Strikethrough on error code cell
        struck = _is_struck(cells[1]) if len(cells) > 1 else False

        errors.append({
            'code': code,
            'key': key,
            'reason': reason,
            'struck': struck,
        })

    # Deduplicate
    seen, deduped = set(), []
    for e in errors:
        k = (e['code'], e['key'])
        if k not in seen:
            seen.add(k)
            deduped.append(e)

    active = [e for e in deduped if not e['struck']]
    struck_list = [e for e in deduped if e['struck']]
    return active, struck_list


# ---------------------------------------------------------------------------
# 5. SOAP TSD detection & parsing
# ---------------------------------------------------------------------------

def _is_soap_tsd(doc) -> bool:
    """Return True if the doc contains SOAP XML paragraphs (soapenv:Envelope)."""
    for p in doc.paragraphs:
        t = p.text.strip()
        if '<soapenv:Envelope' in t:
            return True
    return False


def _extract_soap_xml_from_paragraphs(doc) -> str:
    """Extract the first complete SOAP request Envelope from doc paragraphs."""
    xml_parts = []
    in_envelope = False
    for p in doc.paragraphs:
        line = p.text.strip()
        if '<soapenv:Envelope' in line:
            in_envelope = True
        if in_envelope:
            xml_parts.append(line)
            if '</soapenv:Envelope>' in line:
                break
    return '\n'.join(xml_parts)


def _parse_soap_tsd_fields(doc) -> Tuple[List[dict], List[dict]]:
    """
    Extract request fields from a SOAP TSD docx.
    Strategy 1: mapping table (No | Input/Output | Field Name | Remarks).
    Strategy 2: fall back to leaf nodes in SOAP body XML.
    """
    active_fields: List[dict] = []

    # Strategy 1: mapping table
    mapping_tbl = (_find_table(doc.tables, ['input', 'field'])
                   or _find_table(doc.tables, ['field', 'remarks']))
    if mapping_tbl:
        rows = _table_to_rows(mapping_tbl)
        for row in rows[1:]:          # skip header
            if len(row) < 3:
                continue
            in_out = row[1].lower() if len(row) > 1 else ''
            fname  = row[2].strip()   if len(row) > 2 else ''
            desc   = row[3].strip()   if len(row) > 3 else ''
            if not fname or fname.lower() in ('field name/field', 'field name'):
                continue
            is_input = ('in argument' in in_out or 'input' in in_out
                        or 'selection' in in_out)
            if is_input:
                active_fields.append({
                    'name':        fname,
                    'type':        '',
                    'mandatory':   'Y' if 'mandatory' in desc.lower() else '',
                    'description': desc,
                    'level':       1,
                })

    # Strategy 2: SOAP XML body fallback
    if not active_fields and _HAS_SOAP_UTILS:
        xml_str = _extract_soap_xml_from_paragraphs(doc)
        if xml_str and detect_soap_body(xml_str):
            body_dict = soap_body_to_flat_dict(xml_str)
            seen: set = set()
            for path, val in body_dict.items():
                parts = path.split('.')
                # Skip auth header fields
                if any(p in ('WebRequestCommon', 'Header') for p in parts):
                    continue
                leaf = parts[-1]
                if leaf in seen:
                    continue
                seen.add(leaf)
                active_fields.append({
                    'name':        leaf,
                    'type':        '',
                    'mandatory':   '',
                    'description': f'Sample: {val}' if val else '',
                    'level':       len(parts),
                })

    return active_fields, []


def _parse_soap_tsd_response_fields(doc) -> dict:
    """
    Extract response (output) fields from SOAP TSD mapping table.
    Returns {field_name: type_str} matching the response_data_fields format.
    """
    response_fields: dict = {}

    mapping_tbl = (_find_table(doc.tables, ['input', 'field'])
                   or _find_table(doc.tables, ['field', 'remarks']))
    if not mapping_tbl:
        return response_fields

    rows = _table_to_rows(mapping_tbl)
    for row in rows[1:]:
        if len(row) < 3:
            continue
        in_out = row[1].lower() if len(row) > 1 else ''
        fname  = row[2].strip()  if len(row) > 2 else ''
        if not fname or fname.lower() in ('field name/field', 'field name',
                                          'fields to display (out argument)',
                                          'selection fields (in argument)'):
            continue
        is_output = ('out argument' in in_out or 'output' in in_out
                     or 'display' in in_out)
        if is_output:
            # Infer type from description if possible
            desc = row[3].strip() if len(row) > 3 else ''
            ftype = 'String'
            if 'amount' in fname.lower() or 'balance' in fname.lower():
                ftype = 'Number'
            response_fields[fname] = ftype

    return response_fields


def _parse_soap_tsd_error_table(tables) -> List[dict]:
    """
    Parse SOAP TSD error table where 'Error code' column is empty.
    Creates normalized keys from the 'Error msg' column.
    Returns list of {code, key, reason, struck}.
    """
    tbl = (_find_table(tables, ['error code', 'error msg'])
           or _find_table(tables, ['error code', 'description'])
           or _find_table(tables, ['error msg']))
    if not tbl:
        return []

    errors = []
    rows = _table_to_rows(tbl)
    for row in rows[1:]:   # skip header
        if len(row) < 2:
            continue
        msg  = row[1].strip() if len(row) > 1 else ''
        if not msg:
            continue
        desc = row[2].strip() if len(row) > 2 else ''
        # Normalize message → KEY (UPPER, spaces/dashes/em-dashes → _)
        key = re.sub(r'[^A-Z0-9]+', '_',
                     msg.upper()
                     .replace('\u2013', '_')   # en-dash
                     .replace('\u2014', '_')   # em-dash
                     .replace('-', '_')
                     ).strip('_')
        errors.append({
            'code':   '',
            'key':    key,
            'reason': msg,
            'description': desc,
            'struck': False,
        })
    return errors


def _extract_soap_error_samples(doc) -> list:
    """
    Find 'Sample Error Case' sections in doc paragraphs and extract:
    - empty_fields: XML leaf tag names that are empty in the error request
    - messages_text: human-readable error from SOAP <messages> element
    - normalized_key: normalized form for matching to error table keys
    Returns list of sample dicts.
    """
    _SKIP_TAGS = {
        'envelope', 'header', 'body', 'webrequestcommon',
        'enquiryinputcollection',
    }
    paras = doc.paragraphs
    samples = []
    i = 0
    while i < len(paras):
        text = paras[i].text.strip().lower()
        if 'sample error case' in text or ('error case' in text and 'sample' in text):
            req_parts, resp_parts = [], []
            in_req = in_resp = False
            for j in range(i + 1, min(i + 150, len(paras))):
                t = paras[j].text.strip()
                # Stop at next "Sample ..." section
                if j > i + 5 and 'sample' in t.lower() and 'case' in t.lower():
                    break
                if not in_req and not req_parts and '<soapenv:Envelope' in t:
                    in_req = True
                if in_req:
                    req_parts.append(t)
                    if '</soapenv:Envelope>' in t:
                        in_req = False
                elif req_parts and not resp_parts and '<S:Envelope' in t:
                    in_resp = True
                if in_resp:
                    resp_parts.append(t)
                    if '</S:Envelope>' in t:
                        break

            req_xml  = '\n'.join(req_parts)
            resp_xml = '\n'.join(resp_parts)

            # Find empty leaf elements: <tag></tag>
            empty_fields = []
            if req_xml:
                for m in re.finditer(
                    r'<([A-Za-z][A-Za-z0-9_]*)>\s*</[A-Za-z][A-Za-z0-9_]*>',
                    req_xml
                ):
                    tag = m.group(1)
                    if tag.lower() not in _SKIP_TAGS:
                        empty_fields.append(tag)

            # Extract human-readable error from <messages> in error response
            messages_text = ''
            normalized_key = ''
            if resp_xml:
                m = re.search(r'<messages>(.*?)</messages>', resp_xml, re.DOTALL)
                if m:
                    raw = m.group(1).strip()
                    # Format: "CODE/REF/STATUS/TYPE,HUMAN TEXT" → take after last comma
                    parts = raw.split(',')
                    messages_text = parts[-1].strip() if len(parts) > 1 else raw
                    normalized_key = re.sub(
                        r'[^A-Z0-9]+', '_',
                        messages_text.upper()
                            .replace('\u2013', '_')
                            .replace('\u2014', '_')
                            .replace('-', '_')
                    ).strip('_')

            if empty_fields or messages_text:
                samples.append({
                    'empty_fields':   empty_fields,
                    'messages_text':  messages_text,
                    'normalized_key': normalized_key,
                })
        i += 1
    return samples


def _extract_business_conditions(doc) -> list:
    """
    Extract business validation conditions from doc prose paragraphs.
    Looks for text between the API title and the first XML sample that
    describes validation rules (e.g. field must/must not conditions).
    Returns a list of condition strings.
    """
    conditions = []
    found_title = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if 'API to' in t and not found_title:
            found_title = True
            continue
        if found_title:
            if '<soapenv' in t or '<?xml' in t or 'Sample' in t:
                break
            lower = t.lower()
            # Capture condition-like paragraphs
            if any(kw in lower for kw in [
                'must not', 'must be', 'will validate', 'condition',
                'not equal to', 'greater than', 'less than',
                'will not be', 'will be returned', 'not satisfied',
                'blacklist',
            ]):
                conditions.append(t)
    return conditions


def _extract_soap_operation_slug(doc) -> str:
    """
    Extract the SOAP operation element name from inside <soapenv:Body>.
    E.g. <sup:GetRetailAccountsofCustomer> → GetRetailAccountsofCustomer
    """
    _skip = {'envelope', 'header', 'body', 'webrequestcommon'}
    in_body = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if '<soapenv:Body>' in t:
            in_body = True
            continue
        if '</soapenv:Body>' in t or '</soapenv:Envelope>' in t:
            in_body = False
        if in_body:
            m = re.search(r'<(?:\w+:)?([A-Z][A-Za-z0-9]+)[\s/>]', t)
            if m:
                name = m.group(1)
                if name.lower() not in _skip and not name.endswith('Type'):
                    return name
    return ''


# ---------------------------------------------------------------------------
# 6. Main parse function — returns same contract dict as HTML parser
# ---------------------------------------------------------------------------

def parse_docx_file(docx_path: Path) -> Optional[dict]:
    """
    Parse a single .docx file and return a contract dict with the same
    structure as parse_html_docs produces:

        {
            'method':               str,   # from sampler or '?'
            'doc_method':           str,   # from doc
            'method_is_doc_error':  bool,
            'url':                  str,
            'doc_path':             str,   # e.g. /get-saving-account-transactions
            'active_request_fields': [...],
            'struck_request_fields': [...],
            'active_errors':         [...],
            'struck_errors':         [...],
        }

    The 'method' and 'url' fields are set to doc values (no sampler compare here);
    sampler comparison happens in the main parse_html_docs main block.
    """
    if not HAS_DOCX:
        print(f'  ❌  python-docx chưa cài. Chạy: pip install python-docx')
        return None

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        print(f'  ❌  Không đọc được {docx_path.name}: {e}')
        return None

    tables = doc.tables

    # ── Detect SOAP TSD document ──────────────────────────────────────────────
    if _is_soap_tsd(doc):
        soap_slug = _extract_soap_operation_slug(doc)
        doc_path  = f'/{soap_slug}' if soap_slug else ''
        active_fields, struck_fields = _parse_soap_tsd_fields(doc)
        active_errors, struck_errors = parse_error_table(tables)
        # Fallback: SOAP TSD error tables typically have empty 'Error code' column
        # which causes parse_error_table() to skip all rows — use dedicated parser
        if not active_errors:
            active_errors = _parse_soap_tsd_error_table(tables)
            struck_errors = []
        soap_xml = _extract_soap_xml_from_paragraphs(doc)
        response_data_fields = _parse_soap_tsd_response_fields(doc)
        soap_error_samples   = _extract_soap_error_samples(doc)
        business_conditions  = _extract_business_conditions(doc)
        print(f'  ℹ️  SOAP TSD detected → slug="{soap_slug}", '
              f'{len(active_fields)} request field(s), '
              f'{len(response_data_fields)} response field(s), '
              f'xml_body={bool(soap_xml)}, '
              f'error_samples={len(soap_error_samples)}, '
              f'conditions={len(business_conditions)}')
        return {
            'method':                'POST',
            'doc_method':            'POST',
            'method_is_doc_error':   False,
            'url':                   '?',
            'doc_path':              doc_path,
            'is_soap':               True,
            'sampler_body':          {},
            'sampler_body_format':   'xml',
            'sampler_body_raw_xml':  soap_xml or None,
            'sampler_headers':       {},
            'sampler_prerequest':    [],
            'sampler_setup_items':   [],
            'sampler_extra_variables': {},
            'response_data_fields':  response_data_fields,
            'soap_error_samples':    soap_error_samples,
            'business_conditions':   business_conditions,
            'active_request_fields':  active_fields,
            'struck_request_fields':  [f['name'] for f in struck_fields],
            'active_errors':           active_errors,
            'struck_errors':           [e['code'] for e in struck_errors],
        }
    # ─────────────────────────────────────────────────────────────────────────

    # API Name → doc_path
    api_name = extract_api_name(tables)
    doc_path = api_name if api_name.startswith('/') else f'/{api_name}' if api_name else ''

    # Method
    doc_method = extract_method(tables) or 'NOT FOUND'

    # Request fields
    active_fields, struck_fields = parse_request_table(tables)

    # Error codes
    active_errors, struck_errors = parse_error_table(tables)

    return {
        'method':              doc_method,    # will be overridden by sampler if available
        'doc_method':          doc_method,
        'method_is_doc_error': False,         # will be recalculated with sampler
        'url':                 '?',           # will be filled from sampler
        'doc_path':            doc_path,
        'active_request_fields': active_fields,
        'struck_request_fields': [f['name'] for f in struck_fields],
        'active_errors':          active_errors,
        'struck_errors':          [e['code'] for e in struck_errors],
    }


def slug_from_docx(docx_path: Path) -> str:
    """Derive a slug from a .docx file. Tries:
    1. Parse 'API Name' from doc tables
    2. Filename stem (kebab-to-camelCase)
    """
    if HAS_DOCX:
        try:
            doc = Document(str(docx_path))
            # SOAP TSD: use operation name as slug
            if _is_soap_tsd(doc):
                slug = _extract_soap_operation_slug(doc)
                if slug:
                    return slug
            # Standard REST doc: API Name table
            api_name = extract_api_name(doc.tables)
            if api_name:
                return _kebab_to_slug(api_name)
        except Exception:
            pass

    # Fallback: filename stem
    stem = docx_path.stem
    # Clean common prefixes
    stem = re.sub(r'^(US[\d.+]+|API)\s*[-_]?\s*', '', stem, flags=re.IGNORECASE).strip()
    return _kebab_to_slug(stem) if '-' in stem else stem


def slug_from_doc(doc_path: Path) -> str:
    """Derive a slug from a legacy .doc file (by filename, no parsing)."""
    stem = doc_path.stem
    stem = re.sub(r'^(US[\d.+]+|API)\s*[-_]?\s*', '', stem, flags=re.IGNORECASE).strip()
    return _kebab_to_slug(stem) if '-' in stem else stem


# ---------------------------------------------------------------------------
# CLI: standalone usage
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    import sys
    import json

    ROOT = Path(__file__).resolve().parent.parent
    DOCS = ROOT / 'docs'

    if not HAS_DOCX:
        print('❌  python-docx chưa cài. Chạy: pip install python-docx')
        sys.exit(1)

    # Discover .doc / .docx files
    if len(sys.argv) > 1:
        docx_files = [Path(p) for p in sys.argv[1:]]
    else:
        docx_files = sorted(DOCS.glob('*.docx')) + sorted(DOCS.glob('*.doc'))
        # Remove .doc if a matching .docx already exists
        docx_stems = {p.stem for p in docx_files if p.suffix == '.docx'}
        docx_files = [p for p in docx_files if p.suffix == '.docx' or p.stem not in docx_stems]

    if not docx_files:
        print(f'⚠️  Không tìm thấy file .doc/.docx nào trong {DOCS}')
        sys.exit(0)

    contracts = {}
    for path in docx_files:
        # Convert .doc → .docx if needed
        if path.suffix.lower() == '.doc':
            print(f'\n  🔄  Converting {path.name} → .docx ...')
            converted = convert_doc_to_docx(path)
            if converted is None:
                print(f'  ❌  Bỏ qua {path.name} (conversion failed)')
                continue
            path = converted

        slug = slug_from_docx(path)
        print(f'\n{"="*72}')
        print(f'  {slug}  ←  {path.name}')
        print(f'{"="*72}')

        contract = parse_docx_file(path)
        if contract is None:
            print(f'  ❌  Bỏ qua {path.name}')
            continue

        print(f'  Method:  {contract["doc_method"]}')
        print(f'  Path:    {contract["doc_path"]}')
        print(f'  Fields:  {len(contract["active_request_fields"])} active, '
              f'{len(contract["struck_request_fields"])} struck')
        print(f'  Errors:  {len(contract["active_errors"])} active, '
              f'{len(contract["struck_errors"])} struck')

        contracts[slug] = contract

    if contracts:
        out = ROOT / 'scripts' / 'contracts_from_html.json'
        with open(out, 'w', encoding='utf-8') as f:
            json.dump(contracts, f, ensure_ascii=False, indent=2)
        print(f'\n{"="*72}')
        print(f'  contracts_from_html.json  →  {out}')
        print(f'  {len(contracts)} API(s) parsed from .docx')
        print(f'{"="*72}')
